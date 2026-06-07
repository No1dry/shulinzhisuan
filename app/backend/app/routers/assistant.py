"""
Smart Assistant Router - Backend LLM proxy
Receives chat requests from frontend, queries local database for context,
then forwards to the configured LLM API.
Also supports generating Word reports via LLM.
"""
import os
import json
import uuid
import re
from datetime import datetime
from typing import List, Optional
from fastapi import APIRouter, Depends
from fastapi.responses import FileResponse, StreamingResponse
from sqlalchemy.orm import Session
from sqlalchemy import func
import httpx
import json

from app.database import get_db
from app.models import ResidentMaster, Housing
from app.config import settings
from app.schemas import ResponseModel
from app.table_generator import generate_table, get_table_types
from app.skill_engine import scan_skills, get_skill, get_skill_list

router = APIRouter(prefix="/assistant", tags=["Assistant"])

# Directory for generated reports
REPORTS_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "reports")
os.makedirs(REPORTS_DIR, exist_ok=True)


def _build_system_prompt(db: Session) -> str:
    """Build system prompt with local database context."""
    try:
        total = db.query(ResidentMaster).count()
        if total == 0:
            return (
                "你是「数邻智算」社区治理智能助手，帮助网格员处理日常工作。\n"
                "当前居民数据库为空，请先上传居民总表。\n"
                "你可以回答关于社区治理的一般性问题。\n"
                "回答要简洁、实用，使用中文。"
            )

        key_pop = db.query(ResidentMaster).filter(ResidentMaster.is_key_population == True).count()
        elderly_60 = db.query(ResidentMaster).filter(ResidentMaster.age >= 60).count()
        elderly_alone = db.query(ResidentMaster).filter(
            ResidentMaster.is_living_alone == True, ResidentMaster.age >= 60
        ).count()
        disabled = db.query(ResidentMaster).filter(ResidentMaster.is_disabled == True).count()
        low_income = db.query(ResidentMaster).filter(ResidentMaster.is_low_income == True).count()
        left_behind = db.query(ResidentMaster).filter(ResidentMaster.is_left_behind_child == True).count()

        # Grid breakdown
        grid_stats = []
        grid_names = db.query(ResidentMaster.grid_name).distinct().filter(
            ResidentMaster.grid_name.isnot(None)
        ).all()
        for g in grid_names:
            gn = g[0] if g and g[0] else None
            if not gn:
                continue
            count = db.query(ResidentMaster).filter(ResidentMaster.grid_name == gn).count()
            grid_stats.append(f"  - {gn}: {count}人")

        # Housing stats - wrap in try in case Housing table doesn't exist yet
        try:
            total_houses = db.query(Housing).count()
        except Exception:
            total_houses = 0

        grid_text = "\n".join(grid_stats[:10]) if grid_stats else "  （暂无网格数据）"

        parts = [
            "你是「数邻智算」社区治理智能助手，帮助网格员处理日常工作。",
            "你拥有当前社区的实时居民数据，可以基于这些数据进行查询和统计。\n",
            "【社区概况】",
            f"- 总居民数: {total}人",
            f"- 60岁以上老人: {elderly_60}人",
            f"- 独居老人: {elderly_alone}人",
            f"- 残疾人: {disabled}人",
            f"- 低保户: {low_income}人",
            f"- 留守儿童: {left_behind}人",
            f"- 重点人群: {key_pop}人",
            f"- 房屋数: {total_houses}户\n",
            "【网格分布】",
            f"{grid_text}\n",
            "【工作原则】",
            "1. 回答要简洁、实用，使用中文",
            "2. 涉及具体居民信息时注意隐私保护，只提供脱敏数据",
            "3. 不确定的数据要说明是估算值",
            "4. 可以建议下一步工作措施",
            "5. 如果用户问的数据你拿不准，请说明需要进一步核实",
        ]
        return "\n".join(parts)
    except Exception as e:
        return (
            "你是「数邻智算」社区治理智能助手，帮助网格员处理日常工作。\n"
            f"（数据库查询出错: {str(e)}，将使用通用知识回答）\n"
            "回答要简洁、实用，使用中文。"
        )


def _safe_str(val) -> str:
    """Safely convert any value to string, handling None."""
    if val is None:
        return ""
    return str(val)


def _query_residents_for_context(db: Session, question: str) -> str:
    """Query specific resident data based on the question."""
    try:
        context_parts = []
        q = question.lower()

        if '独居' in q or '老人' in q:
            residents = db.query(ResidentMaster).filter(
                ResidentMaster.is_living_alone == True, ResidentMaster.age >= 60
            ).limit(20).all()
            if residents:
                lines = []
                for r in residents:
                    name = _safe_str(r.name_masked) or "匿名"
                    gender = _safe_str(r.gender) or "未知"
                    age = _safe_str(r.age) or "未知"
                    grid = _safe_str(r.grid_name) or "-"
                    addr = _safe_str(r.residence_address) or "-"
                    lines.append(f"  - {name}, {gender}, {age}岁, 网格:{grid}, 地址:{addr}")
                context_parts.append("【独居老人名单（" + str(len(residents)) + "人）】\n" + "\n".join(lines))

        if '低保' in q:
            count = db.query(ResidentMaster).filter(ResidentMaster.is_low_income == True).count()
            context_parts.append("【低保户统计】共" + str(count) + "人")

        if '残疾' in q:
            residents = db.query(ResidentMaster).filter(ResidentMaster.is_disabled == True).limit(20).all()
            if residents:
                types = {}
                for r in residents:
                    t = _safe_str(r.disability_type) or "未分类"
                    types[t] = types.get(t, 0) + 1
                type_str = ", ".join(f"{k}:{v}人" for k, v in types.items())
                context_parts.append("【残疾人统计】共" + str(len(residents)) + "人，类别分布: " + type_str)

        if '重点' in q:
            residents = db.query(ResidentMaster).filter(ResidentMaster.is_key_population == True).limit(20).all()
            if residents:
                types = {}
                for r in residents:
                    t = _safe_str(r.key_population_type) or "未分类"
                    types[t] = types.get(t, 0) + 1
                type_str = ", ".join(f"{k}:{v}人" for k, v in types.items())
                context_parts.append("【重点人群统计】共" + str(len(residents)) + "人，类别: " + type_str)

        if '网格' in q:
            grids = db.query(ResidentMaster.grid_name, func.count(ResidentMaster.id)).group_by(
                ResidentMaster.grid_name
            ).all()
            if grids:
                lines = []
                for g in grids:
                    gn = g[0] if g and g[0] else None
                    if not gn:
                        continue
                    cnt = g[1] if g[1] is not None else 0
                    lines.append(f"  - {gn}: {cnt}人")
                context_parts.append("【网格人口分布】\n" + "\n".join(lines))

        return "\n\n".join(context_parts) if context_parts else ""
    except Exception as e:
        return ""



@router.post("/chat", response_model=ResponseModel)
async def chat(payload: dict, db: Session = Depends(get_db)):
    """
    Chat with the Smart Assistant.
    Backend receives the question, queries local DB for context,
    then forwards to the configured LLM API.
    """
    question = payload.get("question", "").strip()
    if not question:
        return ResponseModel(code=400, message="请输入问题", data=None)

    # Check LLM configuration
    api_url = settings.LLM_API_URL
    api_key = settings.LLM_API_KEY
    model = settings.LLM_MODEL

    if not api_url or not api_key:
        return ResponseModel(
            code=500,
            message="大模型未配置",
            data={
                "hint": "请在后端 config.py 中配置 LLM_API_URL、LLM_API_KEY 和 LLM_MODEL",
                "example": {
                    "LLM_API_URL": "https://api.moonshot.cn/v1/chat/completions",
                    "LLM_API_KEY": "your-api-key",
                    "LLM_MODEL": "moonshot-v1-8k",
                }
            }
        )

    try:
        # Build system prompt with DB context
        system_prompt = _build_system_prompt(db) or ""

        # Add specific context based on the question
        extra_context = _query_residents_for_context(db, question) or ""
        if extra_context:
            system_prompt = system_prompt + "\n\n【用户问题相关数据】\n" + extra_context

        # Build messages
        messages: List[dict] = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": question},
        ]

        # Forward to LLM API
        async with httpx.AsyncClient(timeout=60.0) as client:
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}",
            }
            body = {
                "model": model,
                "messages": messages,
                "temperature": 0.7,
                "max_tokens": 2048,
            }

            res = await client.post(api_url, headers=headers, json=body)

            if res.status_code != 200:
                return ResponseModel(
                    code=500,
                    message=f"大模型请求失败: HTTP {res.status_code}",
                    data={"detail": res.text[:500]}
                )

            result = res.json()

            # Extract response content (OpenAI-compatible format)
            if "choices" in result and len(result["choices"]) > 0:
                answer = result["choices"][0]["message"]["content"]
            elif "output" in result:
                # Aliyun/Alibaba format
                answer = result["output"]["text"]
            else:
                answer = json.dumps(result, ensure_ascii=False)[:1000]

            return ResponseModel(
                code=200,
                message="success",
                data={
                    "answer": answer,
                    "model": model,
                }
            )

    except httpx.TimeoutException:
        return ResponseModel(code=504, message="大模型请求超时，请稍后重试", data=None)
    except Exception as e:
        return ResponseModel(code=500, message=f"请求失败: {str(e)}", data=None)


def _is_report_request(question: str) -> bool:
    """Check if the user is asking to generate a report."""
    q = question.lower()
    patterns = [
        r'生成.*报告',
        r'生成.*报表',
        r'输出.*报告',
        r'写.*报告',
        r'导出.*报告',
        r'.*的报告',
        r'统计报告',
        r'汇总报告',
        r'分析报告',
    ]
    return any(re.search(p, q) for p in patterns)


def _extract_report_topic(question: str) -> str:
    """Extract the report topic from the question."""
    # Remove common report-related phrases
    cleaned = re.sub(r'(生成|输出|写|导出|一个|份|的|请|帮我|给我)', '', question)
    cleaned = re.sub(r'报告|报表', '', cleaned)
    cleaned = cleaned.strip()
    return cleaned or "社区治理"


def _gather_report_data(db: Session) -> dict:
    """Gather comprehensive data for report generation."""
    data = {"generated_at": datetime.now().strftime("%Y年%m月%d日")}
    
    # Basic stats
    data["total_residents"] = db.query(ResidentMaster).count()
    data["total_grids"] = db.query(ResidentMaster.grid_name).distinct().filter(
        ResidentMaster.grid_name.isnot(None)
    ).count()
    data["total_houses"] = db.query(Housing).count()
    
    # Special groups
    data["elderly_60"] = db.query(ResidentMaster).filter(ResidentMaster.age >= 60).count()
    data["elderly_80"] = db.query(ResidentMaster).filter(ResidentMaster.age >= 80).count()
    data["elderly_alone"] = db.query(ResidentMaster).filter(
        ResidentMaster.is_living_alone == True, ResidentMaster.age >= 60
    ).count()
    data["disabled"] = db.query(ResidentMaster).filter(ResidentMaster.is_disabled == True).count()
    data["low_income"] = db.query(ResidentMaster).filter(ResidentMaster.is_low_income == True).count()
    data["left_behind"] = db.query(ResidentMaster).filter(
        ResidentMaster.is_left_behind_child == True
    ).count()
    data["key_population"] = db.query(ResidentMaster).filter(
        ResidentMaster.is_key_population == True
    ).count()
    
    # Gender distribution
    data["male"] = db.query(ResidentMaster).filter(ResidentMaster.gender == "男").count()
    data["female"] = db.query(ResidentMaster).filter(ResidentMaster.gender == "女").count()
    
    # Grid breakdown
    grid_stats = []
    grid_names = db.query(ResidentMaster.grid_name).distinct().filter(
        ResidentMaster.grid_name.isnot(None)
    ).all()
    for g in grid_names:
        if not g[0]:
            continue
        count = db.query(ResidentMaster).filter(ResidentMaster.grid_name == g[0]).count()
        elderly = db.query(ResidentMaster).filter(
            ResidentMaster.grid_name == g[0], ResidentMaster.age >= 60
        ).count()
        key_pop = db.query(ResidentMaster).filter(
            ResidentMaster.grid_name == g[0], ResidentMaster.is_key_population == True
        ).count()
        grid_stats.append({"name": g[0], "residents": count, "elderly": elderly, "key_pop": key_pop})
    data["grid_stats"] = sorted(grid_stats, key=lambda x: x["residents"], reverse=True)
    
    return data


def _strip_markdown(text: str) -> str:
    """Remove markdown formatting for Word document."""
    if not text:
        return ""
    # Remove **bold** and *italic*
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    # Remove # headings
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    # Remove `code`
    text = re.sub(r'`(.+?)`', r'\1', text)
    # Remove > blockquote
    text = re.sub(r'^>\s+', '', text, flags=re.MULTILINE)
    # Remove --- horizontal rules
    text = re.sub(r'^---+$', '', text, flags=re.MULTILINE)
    return text.strip()


def _set_run_font(run, font_name='宋体', font_size=None, bold=False, color=None):
    """Set font properties for a run."""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    if font_size is None:
        font_size = Pt(12)
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    # For Chinese font support
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def _create_word_report(title: str, content: str, data: dict) -> str:
    """Create a Word document report with Song font and clean markdown. Returns file path."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        
        doc = Document()
        
        # Set default font for Normal style
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # Title
        title_para = doc.add_heading(level=0)
        title_run = title_para.add_run(title)
        _set_run_font(title_run, '宋体', Pt(22), bold=True)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle with date
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_para.add_run(f"生成日期：{data['generated_at']}")
        _set_run_font(subtitle_run, '宋体', Pt(12), color=RGBColor(128, 128, 128))
        doc.add_paragraph()
        
        # Summary stats table
        heading_para = doc.add_paragraph()
        heading_run = heading_para.add_run('一、社区概况')
        _set_run_font(heading_run, '宋体', Pt(16), bold=True)
        
        stats_table = doc.add_table(rows=1, cols=2)
        stats_table.style = 'Table Grid'
        hdr = stats_table.rows[0].cells
        hdr[0].text = '指标'
        hdr[1].text = '数值'
        
        stat_rows = [
            ("居民总数", f"{data['total_residents']} 人"),
            ("网格数量", f"{data['total_grids']} 个"),
            ("房屋数量", f"{data['total_houses']} 户"),
            ("60岁以上老人", f"{data['elderly_60']} 人"),
            ("80岁以上高龄老人", f"{data['elderly_80']} 人"),
            ("独居老人", f"{data['elderly_alone']} 人"),
            ("残疾人", f"{data['disabled']} 人"),
            ("低保户", f"{data['low_income']} 人"),
            ("留守儿童", f"{data['left_behind']} 人"),
            ("重点人群", f"{data['key_population']} 人"),
            ("男性", f"{data['male']} 人"),
            ("女性", f"{data['female']} 人"),
        ]
        for label, value in stat_rows:
            row = stats_table.add_row().cells
            row[0].text = label
            row[1].text = value
        
        doc.add_paragraph()
        
        # Grid breakdown
        if data["grid_stats"]:
            heading_para = doc.add_paragraph()
            heading_run = heading_para.add_run('二、网格分布')
            _set_run_font(heading_run, '宋体', Pt(16), bold=True)
            
            grid_table = doc.add_table(rows=1, cols=4)
            grid_table.style = 'Table Grid'
            hdr = grid_table.rows[0].cells
            hdr[0].text = '网格名称'
            hdr[1].text = '居民人数'
            hdr[2].text = '老人数'
            hdr[3].text = '重点人群'
            
            for g in data["grid_stats"]:
                row = grid_table.add_row().cells
                row[0].text = g["name"]
                row[1].text = str(g["residents"])
                row[2].text = str(g["elderly"])
                row[3].text = str(g["key_pop"])
            
            doc.add_paragraph()
        
        # LLM-generated content with markdown stripped
        heading_para = doc.add_paragraph()
        heading_run = heading_para.add_run('三、详细分析')
        _set_run_font(heading_run, '宋体', Pt(16), bold=True)
        
        clean_content = _strip_markdown(content)
        for para_text in clean_content.split('\n\n'):
            stripped = para_text.strip()
            if stripped:
                # Handle lines that start with ## as sub-headings
                if stripped.startswith('## '):
                    sub_para = doc.add_paragraph()
                    sub_run = sub_para.add_run(_strip_markdown(stripped))
                    _set_run_font(sub_run, '宋体', Pt(14), bold=True)
                elif stripped.startswith('- ') or stripped.startswith('* '):
                    # Bullet points
                    for line in stripped.split('\n'):
                        line = line.strip()
                        if line.startswith('- ') or line.startswith('* '):
                            line = line[2:]
                        if line:
                            p = doc.add_paragraph(style='List Bullet')
                            run = p.add_run(_strip_markdown(line))
                            _set_run_font(run, '宋体', Pt(12))
                else:
                    p = doc.add_paragraph()
                    run = p.add_run(_strip_markdown(stripped))
                    _set_run_font(run, '宋体', Pt(12))
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run("本报告由数邻智算系统自动生成")
        _set_run_font(footer_run, '宋体', Pt(9), color=RGBColor(128, 128, 128))
        
        # Save
        filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:6]}.docx"
        file_path = os.path.join(REPORTS_DIR, filename)
        doc.save(file_path)
        return file_path
        
    except ImportError:
        # Fallback: create markdown file
        filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:6]}.md"
        file_path = os.path.join(REPORTS_DIR, filename)
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(f"# {title}\n\n")
            f.write(f"生成日期：{data['generated_at']}\n\n")
            f.write("## 社区概况\n\n")
            for label, value in [
                ("居民总数", f"{data['total_residents']} 人"),
                ("60岁以上老人", f"{data['elderly_60']} 人"),
                ("独居老人", f"{data['elderly_alone']} 人"),
                ("残疾人", f"{data['disabled']} 人"),
                ("低保户", f"{data['low_income']} 人"),
                ("重点人群", f"{data['key_population']} 人"),
            ]:
                f.write(f"- {label}：{value}\n")
            f.write(f"\n## 详细分析\n\n{content}\n")
        return file_path


@router.post("/generate-report", response_model=ResponseModel)
async def generate_report(payload: dict, db: Session = Depends(get_db)):
    """
    Generate a Word report via LLM.
    1. Gather DB data
    2. Ask LLM to write report content
    3. Create Word document
    4. Return download URL
    """
    question = payload.get("question", "").strip()
    if not question:
        return ResponseModel(code=400, message="请输入报告要求", data=None)

    api_url = settings.LLM_API_URL
    api_key = settings.LLM_API_KEY
    model = settings.LLM_MODEL

    if not api_url or not api_key:
        return ResponseModel(
            code=500, message="大模型未配置，无法生成报告",
            data={"hint": "请先配置 LLM_API_URL、LLM_API_KEY 和 LLM_MODEL"}
        )

    try:
        # Gather data
        report_data = _gather_report_data(db)
        topic = _extract_report_topic(question)
        
        if report_data["total_residents"] == 0:
            return ResponseModel(code=400, message="居民数据库为空，请先上传居民总表", data=None)

        # Build report-specific system prompt
        data_text = json.dumps(report_data, ensure_ascii=False, indent=2)
        system_prompt = (
            f"你是一位专业的社区治理报告撰写专家。请根据以下数据撰写一份正式的社区治理报告。\n\n"
            f"【报告主题】{topic}\n\n"
            f"【数据】\n{data_text}\n\n"
            f"【写作要求】\n"
            f"1. 报告标题要正式、规范\n"
            f"2. 使用正式公文体\n"
            f"3. 包含数据分析和趋势判断\n"
            f"4. 提出下一步工作建议\n"
            f"5. 适当分节，结构清晰\n"
            f"6. 总字数不少于800字"
        )

        # Call LLM
        async with httpx.AsyncClient(timeout=60.0) as client:
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}",
            }
            body = {
                "model": model,
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"请撰写报告：{topic}"},
                ],
                "temperature": 0.7,
                "max_tokens": 4096,
            }

            res = await client.post(api_url, headers=headers, json=body)
            if res.status_code != 200:
                return ResponseModel(code=500, message=f"大模型请求失败: HTTP {res.status_code}", data=None)

            result = res.json()
            if "choices" in result and len(result["choices"]) > 0:
                report_content = result["choices"][0]["message"]["content"]
            elif "output" in result:
                report_content = result["output"]["text"]
            else:
                report_content = "报告生成失败"

        # Create Word document
        title = f"{topic}报告" if "报告" not in topic else topic
        file_path = _create_word_report(title, report_content, report_data)
        filename = os.path.basename(file_path)

        return ResponseModel(
            code=200,
            message="报告生成成功",
            data={
                "answer": report_content,
                "is_report": True,
                "title": title,
                "filename": filename,
                "download_url": f"/api/assistant/download-report/{filename}",
            }
        )

    except Exception as e:
        return ResponseModel(code=500, message=f"报告生成失败: {str(e)}", data=None)


@router.get("/download-report/{filename}")
async def download_report(filename: str):
    """Download a generated report or table file (.docx or .xlsx)."""
    file_path = os.path.join(REPORTS_DIR, filename)
    if not os.path.exists(file_path):
        return ResponseModel(code=404, message="文件不存在", data=None)

    # Determine correct MIME type based on extension
    if filename.endswith('.xlsx'):
        media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    elif filename.endswith('.docx'):
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif filename.endswith('.md'):
        media_type = "text/markdown"
    else:
        media_type = "application/octet-stream"

    return FileResponse(
        path=file_path,
        filename=filename,
        media_type=media_type
    )


@router.get("/config-status", response_model=ResponseModel)
async def get_config_status():
    """Check if LLM is configured."""
    return ResponseModel(
        code=200,
        message="success",
        data={
            "configured": bool(settings.LLM_API_URL and settings.LLM_API_KEY),
            "model": settings.LLM_MODEL or "",
            "api_url": settings.LLM_API_URL or "",
        }
    )


# ── SSE 流式聊天 ────────────────────────────────

@router.post("/chat-stream")
async def chat_stream(payload: dict, db: Session = Depends(get_db)):
    """
    Stream chat response via SSE.
    Frontend uses EventSource to receive chunks in real-time.
    """
    question = payload.get("question", "").strip()
    if not question:
        async def error_stream():
            yield "data: " + json.dumps({"type": "error", "content": "请输入问题"}, ensure_ascii=False) + "\n\n"
        return StreamingResponse(error_stream(), media_type="text/event-stream")

    api_url = settings.LLM_API_URL
    api_key = settings.LLM_API_KEY
    model = settings.LLM_MODEL

    if not api_url or not api_key:
        async def error_stream():
            yield "data: " + json.dumps({"type": "error", "content": "大模型未配置"}, ensure_ascii=False) + "\n\n"
        return StreamingResponse(error_stream(), media_type="text/event-stream")

    # Build system prompt with DB context
    system_prompt = _build_system_prompt(db) or ""
    extra_context = _query_residents_for_context(db, question) or ""
    if extra_context:
        system_prompt = system_prompt + "\n\n【用户问题相关数据】\n" + extra_context

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": question},
    ]

    async def sse_stream():
        """Generate SSE stream from LLM response."""
        full_content = ""
        try:
            # Check if API supports streaming
            supports_stream = True

            async with httpx.AsyncClient(timeout=60.0) as client:
                headers = {
                    "Content-Type": "application/json",
                    "Authorization": f"Bearer {api_key}",
                }
                body = {
                    "model": model,
                    "messages": messages,
                    "temperature": 0.7,
                    "max_tokens": 2048,
                    "stream": True,
                }

                async with client.stream("POST", api_url, headers=headers, json=body) as res:
                    if res.status_code != 200:
                        error_text = ""
                        async for chunk in res.aiter_text():
                            if chunk:
                                error_text = error_text + chunk
                        yield "data: " + json.dumps({"type": "error", "content": f"大模型请求失败: {res.status_code}"}, ensure_ascii=False) + "\n\n"
                        return

                    async for line in res.aiter_lines():
                        if not line or line.strip() == "":
                            continue
                        if line.startswith("data: "):
                            data_str = line[6:]
                            if data_str == "[DONE]":
                                break
                            try:
                                data = json.loads(data_str)
                                # OpenAI format
                                if "choices" in data and len(data["choices"]) > 0:
                                    delta = data["choices"][0].get("delta", {})
                                    chunk = delta.get("content")
                                    if chunk is not None:
                                        full_content = full_content + chunk
                                        yield "data: " + json.dumps({"type": "chunk", "content": chunk}, ensure_ascii=False) + "\n\n"
                            except (json.JSONDecodeError, TypeError):
                                continue

            # Send done event with full content
            yield "data: " + json.dumps({"type": "done", "content": full_content}, ensure_ascii=False) + "\n\n"

        except Exception as e:
            yield "data: " + json.dumps({"type": "error", "content": str(e)}, ensure_ascii=False) + "\n\n"

    return StreamingResponse(sse_stream(), media_type="text/event-stream")


# ── 表格生成 ────────────────────────────────────

@router.get("/table-types", response_model=ResponseModel)
async def list_table_types():
    """Get available table types for auto-generation."""
    return ResponseModel(code=200, message="success", data={"types": get_table_types()})


@router.post("/generate-table", response_model=ResponseModel)
async def generate_table_endpoint(payload: dict, db: Session = Depends(get_db)):
    """Generate an Excel table of the specified type."""
    table_type = payload.get("table_type", "").strip()
    if not table_type:
        return ResponseModel(code=400, message="请选择表格类型", data=None)

    try:
        file_path = generate_table(table_type, db)
        filename = os.path.basename(file_path)
        return ResponseModel(code=200, message="表格生成成功", data={
            "filename": filename,
            "download_url": f"/api/assistant/download-report/{filename}",
        })
    except ValueError as e:
        return ResponseModel(code=400, message=str(e), data=None)
    except Exception as e:
        return ResponseModel(code=500, message=f"表格生成失败: {str(e)}", data=None)


# ── Excel Maker Skill ───────────────────────────

SKILL_DIR = os.path.join(os.path.dirname(__file__), "..", "skills")


def _read_skill_file(name: str) -> str:
    """Read a skill markdown file."""
    path = os.path.join(SKILL_DIR, f"{name}.md")
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    return ""


def _extract_table_data(file_path: str) -> dict:
    """Read Excel file and extract data for frontend preview."""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(file_path, read_only=True, data_only=True)
        ws = wb.active

        rows = list(ws.iter_rows(values_only=True))
        if len(rows) < 3:  # title + subtitle + header
            return {"headers": [], "rows": []}

        # Skip title rows, find header row (the one with blue background or just the first non-empty)
        header_row_idx = 2  # Usually row index 2 (0-based) after title+subtitle
        headers = [str(h) if h is not None else "" for h in rows[header_row_idx]]
        # Remove trailing empty headers
        while headers and headers[-1] == "":
            headers.pop()

        data_rows = []
        for row in rows[header_row_idx + 1:]:
            row_data = [str(cell) if cell is not None else "" for cell in row[:len(headers)]]
            if any(row_data):  # Skip empty rows
                data_rows.append(row_data)

        wb.close()
        return {"headers": headers, "rows": data_rows}
    except Exception:
        return {"headers": [], "rows": []}


@router.post("/excel-maker", response_model=ResponseModel)
async def excel_maker(payload: dict, db: Session = Depends(get_db)):
    """
    Excel Maker Skill - Generate table based on natural language description.
    Reads skill file, uses LLM to parse demand, generates table with preview.
    """
    demand = payload.get("demand", "").strip()
    if not demand:
        return ResponseModel(code=400, message="请输入表格需求描述", data=None)

    api_url = settings.LLM_API_URL
    api_key = settings.LLM_API_KEY
    model = settings.LLM_MODEL

    if not api_url or not api_key:
        return ResponseModel(code=500, message="大模型未配置", data=None)

    try:
        # Read skill file
        skill_content = _read_skill_file("excel-maker")

        # Build prompt for LLM to parse table type
        system_prompt = (
            "You are an Excel table generation expert. Based on the skill guide below, "
            "determine which table type the user wants and generate a brief explanation.\n\n"
            f"{skill_content}\n\n"
            "Respond ONLY with a JSON object in this exact format:\n"
            '{"table_type": "resident_register", "table_title": "居民信息登记表", "explanation": "已为您生成居民信息登记表..."}'
        )

        user_prompt = f'用户需求："{demand}"\n\n请确定表格类型并生成说明。'

        # Call LLM
        async with httpx.AsyncClient(timeout=30.0) as client:
            res = await client.post(api_url, headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}",
            }, json={
                "model": model,
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                "temperature": 0.3,
                "max_tokens": 512,
            })

            if res.status_code != 200:
                # Fallback: try keyword matching
                table_type = _match_table_type_by_keyword(demand)
                return _generate_table_response(table_type, demand, db)

            result = res.json()
            if "choices" in result and len(result["choices"]) > 0:
                llm_response = result["choices"][0]["message"]["content"]
            else:
                llm_response = "{}"

            # Parse JSON from LLM response
            try:
                # Extract JSON from markdown code blocks if present
                json_match = re.search(r'\{.*\}', llm_response, re.DOTALL)
                if json_match:
                    parsed = json.loads(json_match.group())
                    table_type = parsed.get("table_type", "resident_register")
                else:
                    table_type = _match_table_type_by_keyword(demand)
            except json.JSONDecodeError:
                table_type = _match_table_type_by_keyword(demand)

        return _generate_table_response(table_type, demand, db)

    except Exception as e:
        return ResponseModel(code=500, message=f"生成失败: {str(e)}", data=None)


def _match_table_type_by_keyword(demand: str) -> str:
    """Fallback: match table type by keywords in demand."""
    d = demand.lower()
    if "老人" in d or "独居" in d or "走访" in d:
        return "elderly_visit"
    if "低保" in d:
        return "low_income"
    if "残疾" in d:
        return "disabled_register"
    if "重点" in d:
        return "key_population"
    if "网格" in d or "统计" in d or "汇总" in d:
        return "grid_statistics"
    if "房屋" in d or "楼栋" in d or "住房" in d:
        return "housing_register"
    if "安全" in d or "巡查" in d:
        return "safety_inspection"
    if "诉求" in d or "上访" in d:
        return "appeal_register"
    if "协商" in d or "会议" in d:
        return "negotiation_record"
    return "resident_register"


def _generate_table_response(table_type: str, demand: str, db: Session, custom_title: str = ""):
    """Generate table and build response with preview data. Supports any table type."""
    from app.table_generator import _get_schema

    schema = _get_schema(table_type, custom_title)

    # Generate Excel file (unknown types automatically get generic template)
    file_path = generate_table(table_type, db, custom_title)
    filename = os.path.basename(file_path)

    # Extract preview data
    table_data = _extract_table_data(file_path)

    # Build explanation
    row_count = len(table_data.get("rows", []))
    has_data = row_count > 0
    is_custom = table_type not in [
        "resident_register", "elderly_visit", "low_income", "disabled_register",
        "key_population", "grid_statistics", "housing_register", "safety_inspection",
        "appeal_register", "negotiation_record"
    ]

    explanation = f"已为您生成 **{schema['name']}**。\n\n"
    if is_custom:
        explanation = explanation + "- 这是一个自定义表格，已根据您的需求生成通用模板\n"
    explanation = explanation + f"- 数据行数：{row_count} 行\n- 数据来源：数邻智算居民数据库\n"
    if has_data:
        explanation = explanation + "- 已自动填入居民信息，空白列为需要手动补充的内容\n"
    else:
        explanation = explanation + "- 当前数据库暂无相关数据，已生成空白模板\n"

    return ResponseModel(code=200, message="success", data={
        "table_type": table_type,
        "table_title": schema["name"],
        "explanation": explanation,
        "table_data": table_data,
        "row_count": row_count,
        "download_url": f"/api/assistant/download-report/{filename}",
        "filename": filename,
    })


# ── Skill System ────────────────────────────────

@router.get("/skills", response_model=ResponseModel)
async def list_skills():
    """List all available skills/experts."""
    return ResponseModel(code=200, message="success", data={"skills": get_skill_list()})


@router.post("/skills/{skill_name}", response_model=ResponseModel)
async def execute_skill(skill_name: str, payload: dict, db: Session = Depends(get_db)):
    """
    Execute a skill by name.
    skill_name: excel-maker, report-writer, chat, etc.
    payload: { "demand": "用户需求描述", "question": "原始问题" }
    """
    demand = payload.get("demand", "").strip()
    question = payload.get("question", demand).strip()

    # Route to specific skill handler
    if skill_name == "excel-maker":
        return await _handle_excel_maker(demand, db)
    elif skill_name == "report-writer":
        return await _handle_report_writer(demand, db)
    elif skill_name == "chat":
        return await _handle_chat(question, db)
    else:
        return ResponseModel(code=400, message=f"未知技能: {skill_name}", data=None)


async def _handle_excel_maker(demand: str, db: Session):
    """Handle excel-maker skill."""
    api_url = settings.LLM_API_URL
    api_key = settings.LLM_API_KEY
    model = settings.LLM_MODEL

    if not api_url or not api_key:
        return ResponseModel(code=500, message="大模型未配置", data=None)

    try:
        # Read skill file for prompt
        skill = get_skill("excel-maker")
        skill_content = skill.md_content if skill else ""

        system_prompt = (
            "You are an Excel table generation expert. Based on the skill guide below, "
            "determine which table type the user wants.\n\n"
            f"{skill_content}\n\n"
            "Respond ONLY with a JSON object:\n"
            '{"table_type": "resident_register", "table_title": "居民信息登记表", "explanation": "..."}'
        )

        async with httpx.AsyncClient(timeout=30.0) as client:
            res = await client.post(api_url, headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}",
            }, json={
                "model": model,
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f'用户需求："{demand}"'},
                ],
                "temperature": 0.3,
                "max_tokens": 512,
            })

            table_type = _match_table_type_by_keyword(demand)

            if res.status_code == 200:
                result = res.json()
                if "choices" in result and len(result["choices"]) > 0:
                    llm_response = result["choices"][0]["message"]["content"]
                    try:
                        json_match = re.search(r'\{.*\}', llm_response, re.DOTALL)
                        if json_match:
                            parsed = json.loads(json_match.group())
                            table_type = parsed.get("table_type", table_type)
                    except json.JSONDecodeError:
                        pass

        return _generate_table_response(table_type, demand, db)

    except Exception as e:
        return ResponseModel(code=500, message=f"生成失败: {str(e)}", data=None)


async def _handle_report_writer(demand: str, db: Session):
    """Handle report-writer skill."""
    api_url = settings.LLM_API_URL
    api_key = settings.LLM_API_KEY
    model = settings.LLM_MODEL

    if not api_url or not api_key:
        return ResponseModel(code=500, message="大模型未配置", data=None)

    try:
        # Gather report data
        report_data = _gather_report_data(db)

        if report_data["total_residents"] == 0:
            return ResponseModel(code=400, message="居民数据库为空，请先上传居民总表", data=None)

        topic = _extract_report_topic(demand) if demand else "社区治理工作报告"
        data_text = json.dumps(report_data, ensure_ascii=False, indent=2)

        # Read skill file for prompt
        skill = get_skill("report-writer")
        skill_prompt = skill.system_prompt if skill else ""

        system_prompt = (
            f"{skill_prompt}\n\n"
            f"你是一位专业的社区治理报告撰写专家。请根据以下数据撰写一份正式的报告。\n\n"
            f"【报告主题】{topic}\n\n"
            f"【数据】\n{data_text}\n\n"
            f"【写作要求】\n"
            f"1. 报告标题要正式、规范\n"
            f"2. 使用正式公文体\n"
            f"3. 包含数据分析和趋势判断\n"
            f"4. 提出下一步工作建议\n"
            f"5. 适当分节，结构清晰\n"
            f"6. 总字数不少于800字"
        )

        async with httpx.AsyncClient(timeout=60.0) as client:
            res = await client.post(api_url, headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}",
            }, json={
                "model": model,
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"请撰写报告：{topic}"},
                ],
                "temperature": 0.7,
                "max_tokens": 4096,
            })

            if res.status_code != 200:
                return ResponseModel(code=500, message=f"大模型请求失败: {res.status_code}", data=None)

            result = res.json()
            if "choices" in result and len(result["choices"]) > 0:
                report_content = result["choices"][0]["message"]["content"]
            else:
                report_content = "报告生成失败"

        title = f"{topic}报告" if "报告" not in topic else topic
        file_path = _create_word_report(title, report_content, report_data)
        filename = os.path.basename(file_path)

        return ResponseModel(code=200, message="报告生成成功", data={
            "answer": report_content,
            "is_report": True,
            "title": title,
            "filename": filename,
            "download_url": f"/api/assistant/download-report/{filename}",
            "explanation": f"已为您生成 **{title}**，可点击下方按钮下载 Word 文档。",
        })

    except Exception as e:
        return ResponseModel(code=500, message=f"报告生成失败: {str(e)}", data=None)


async def _handle_chat(question: str, db: Session):
    """Handle chat skill - normal Q&A."""
    # This is handled by the chat-stream endpoint, but we provide a non-streaming fallback
    api_url = settings.LLM_API_URL
    api_key = settings.LLM_API_KEY
    model = settings.LLM_MODEL

    if not api_url or not api_key:
        return ResponseModel(code=500, message="大模型未配置", data=None)

    try:
        system_prompt = _build_system_prompt(db)
        extra_context = _query_residents_for_context(db, question)
        if extra_context:
            system_prompt = system_prompt + "\n\n【用户问题相关数据】\n" + extra_context

        async with httpx.AsyncClient(timeout=60.0) as client:
            res = await client.post(api_url, headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}",
            }, json={
                "model": model,
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": question},
                ],
                "temperature": 0.7,
                "max_tokens": 2048,
            })

            if res.status_code != 200:
                return ResponseModel(code=500, message=f"大模型请求失败: {res.status_code}", data=None)

            result = res.json()
            if "choices" in result and len(result["choices"]) > 0:
                answer = result["choices"][0]["message"]["content"]
            else:
                answer = "请求失败"

            return ResponseModel(code=200, message="success", data={"answer": answer})

    except Exception as e:
        return ResponseModel(code=500, message=f"请求失败: {str(e)}", data=None)